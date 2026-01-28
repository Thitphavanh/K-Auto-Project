from django.shortcuts import render, redirect, get_object_or_404
from django.contrib import messages
from django.contrib.auth.decorators import login_required
from django.views.decorators.http import require_http_methods
from django.http import HttpResponse
from django.db import models
from django.db.models import Sum, Count, Q
from django.db.models.functions import TruncDay, TruncWeek, TruncMonth, TruncYear
from django.utils import timezone
from .models import CashBook
import csv
import datetime
import openpyxl
from openpyxl.styles import Font, Alignment
from docx import Document
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter, A4
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from io import BytesIO
from urllib.parse import quote

@login_required
def dashboard(request):
    """
    Dashboard for CashBook: Lists recent transactions and shows summary.
    """
    today = timezone.now().date()
    
    # Filter by date range (default to today, or allow user to pick)
    date_str = request.GET.get('date')
    if date_str:
        try:
            view_date = timezone.datetime.strptime(date_str, '%Y-%m-%d').date()
        except ValueError:
            view_date = today
    else:
        view_date = today

    # Get transactions for the selected day
    transactions = CashBook.objects.filter(date=view_date).order_by('-time', '-created_at')
    
    # Helper function to calculate balance for a specific currency
    def get_totals(currency_code):
        income = transactions.filter(transaction_type='IN', currency=currency_code).aggregate(Sum('amount'))['amount__sum'] or 0
        expense = transactions.filter(transaction_type='OUT', currency=currency_code).aggregate(Sum('amount'))['amount__sum'] or 0
        balance = income - expense
        
        # All time
        all_income = CashBook.objects.filter(transaction_type='IN', currency=currency_code).aggregate(Sum('amount'))['amount__sum'] or 0
        all_expense = CashBook.objects.filter(transaction_type='OUT', currency=currency_code).aggregate(Sum('amount'))['amount__sum'] or 0
        total_balance = all_income - all_expense
        
        return {
            'income': income,
            'expense': expense,
            'balance': balance,
            'total_balance': total_balance
        }

    # Calculate totals for each currency
    totals = {
        'LAK': get_totals('LAK'),
        'THB': get_totals('THB'),
        'USD': get_totals('USD'),
    }

    # Calculate totals
    lak_stats = get_totals('LAK')
    thb_stats = get_totals('THB')
    usd_stats = get_totals('USD')

    # Prepare WhatsApp String
    wa_date = view_date.strftime('%d/%m/%Y')
    wa_text = f"*📊 ສະຫຼຸບບັນຊີ K-Auto*%0A"
    wa_text += f"📅 ວັນທີ: {wa_date}%0A"
    wa_text += "━━━━━━━━━━━━━━━%0A%0A"

    has_activity = False

    if lak_stats['income'] > 0 or lak_stats['expense'] > 0:
        has_activity = True
        balance_icon = "✅" if lak_stats['balance'] >= 0 else "❌"
        wa_text += f"*💰 LAK (ກີບ)*%0A"
        wa_text += f"   🟢 ລາຍຮັບ: *{lak_stats['income']:,.0f}* ₭%0A"
        wa_text += f"   🔴 ລາຍຈ່າຍ: *{lak_stats['expense']:,.0f}* ₭%0A"
        wa_text += f"   {balance_icon} ຍອດເຫຼືອ: *{lak_stats['balance']:,.0f}* ₭%0A%0A"

    if thb_stats['income'] > 0 or thb_stats['expense'] > 0:
        has_activity = True
        balance_icon = "✅" if thb_stats['balance'] >= 0 else "❌"
        wa_text += f"*💰 THB (ບາດ)*%0A"
        wa_text += f"   🟢 ລາຍຮັບ: *{thb_stats['income']:,.2f}* ฿%0A"
        wa_text += f"   🔴 ລາຍຈ່າຍ: *{thb_stats['expense']:,.2f}* ฿%0A"
        wa_text += f"   {balance_icon} ຍອດເຫຼືອ: *{thb_stats['balance']:,.2f}* ฿%0A%0A"

    if usd_stats['income'] > 0 or usd_stats['expense'] > 0:
        has_activity = True
        balance_icon = "✅" if usd_stats['balance'] >= 0 else "❌"
        wa_text += f"*💰 USD (ໂດລາ)*%0A"
        wa_text += f"   🟢 ລາຍຮັບ: *{usd_stats['income']:,.2f}* $%0A"
        wa_text += f"   🔴 ລາຍຈ່າຍ: *{usd_stats['expense']:,.2f}* $%0A"
        wa_text += f"   {balance_icon} ຍອດເຫຼືອ: *{usd_stats['balance']:,.2f}* $%0A%0A"

    # If no activity
    if not has_activity:
        wa_text += "❌ _ບໍ່ມີການເຄື່ອນໄຫວໃນມື້ນີ້_%0A%0A"

    wa_text += "━━━━━━━━━━━━━━━%0A"
    wa_text += "🏢 *K-Auto Service*"

    context = {
        'view_date': view_date,
        'transactions': transactions,
        'lak': lak_stats,
        'thb': thb_stats,
        'usd': usd_stats,
        'whatsapp_text': wa_text, # Pass to template
    }
    return render(request, 'accounting/dashboard.html', context)

@login_required
def add_transaction(request):
    """
    Form to add a new transaction (Income or Expense).
    """
    if request.method == 'POST':
        date = request.POST.get('date')
        time = request.POST.get('time')
        description = request.POST.get('description')
        transaction_type = request.POST.get('transaction_type')
        currency = request.POST.get('currency')
        amount = request.POST.get('amount')
        category = request.POST.get('category')

        # Validation
        if not description or not amount or not transaction_type or not currency:
            messages.error(request, "ກະລຸນາປ້ອນຂໍ້ມູນໃຫ້ຄົບຖ້ວນ")
            return redirect('accounting:add_transaction')

        try:
            CashBook.objects.create(
                date=date,
                time=time if time else timezone.now().time(),
                description=description,
                transaction_type=transaction_type,
                currency=currency,
                amount=amount,
                category=category
            )
            messages.success(request, f"ບັນທຶກລາຍຮັບ-ລາຍຈ່າຍ ({currency}) ສຳເລັດ")
            return redirect('accounting:dashboard')
        except Exception as e:
            messages.error(request, f"ເກີດຂໍ້ຜິດພາດ: {str(e)}")
            return redirect('accounting:add_transaction')

    # Get all product categories from store app
    from store.models import Category
    categories = Category.objects.all().order_by('name')

    context = {
        'today': timezone.now().date(),
        'categories': categories,
    }
    return render(request, 'accounting/form.html', context)

@login_required
def report_view(request):
    """
    View for generating Accounting Reports (Daily, Weekly, Monthly, Yearly).
    Integrates data from:
    1. CashBook (Manual Income/Expense)
    2. Store Orders (Sales Income)
    3. Store Transactions (Stock/Purchase Expenses - Assumed THB)
    """
def _get_report_data(report_type, start_date_str=None, end_date_str=None):
    from store.models import Order, Transaction # Import here to avoid circular dependencies if any
    
    # Determine truncation function based on report type
    if report_type == 'weekly':
        trunc_func = TruncWeek
    elif report_type == 'monthly':
        trunc_func = TruncMonth
    elif report_type == 'yearly':
        trunc_func = TruncYear
    else: # daily
        trunc_func = TruncDay
        
    # Base Querysets
    cashbook_qs = CashBook.objects.all()
    order_qs = Order.objects.all()
    transaction_qs = Transaction.objects.filter(transaction_type='IN') # Only Stock In (Expenses)
    
    if start_date_str:
        cashbook_qs = cashbook_qs.filter(date__gte=start_date_str)
        order_qs = order_qs.filter(date__gte=start_date_str)
        transaction_qs = transaction_qs.filter(created_at__date__gte=start_date_str)
        
    if end_date_str:
        cashbook_qs = cashbook_qs.filter(date__lte=end_date_str)
        order_qs = order_qs.filter(date__lte=end_date_str)
        transaction_qs = transaction_qs.filter(created_at__date__lte=end_date_str)
        
    # --- 1. Aggregating CashBook ---
    cashbook_data = (
        cashbook_qs
        .annotate(period=trunc_func('date'))
        .values('period', 'currency')
        .annotate(
            total_income=Sum('amount', filter=Q(transaction_type='IN'), default=0),
            total_expense=Sum('amount', filter=Q(transaction_type='OUT'), default=0)
        )
    )
    
    # --- 2. Aggregating Orders (Sales) ---
    order_data = (
        order_qs
        .annotate(period=trunc_func('date'))
        .values('period')
        .annotate(
            income_lak=Sum('net_amount_lak', default=0),
            income_thb=Sum('net_amount_thb', default=0),
            income_usd=Sum('net_amount_usd', default=0)
        )
    )
    
    # --- 3. Aggregating Transactions (Stock Expenses) ---
    stock_data = (
        transaction_qs
        .annotate(period=trunc_func('created_at'))
        .values('period')
        .annotate(
            expense_thb=Sum('total_value', default=0)
        )
    )
    
    # --- 4. Merging Data Python-side ---
    grouped_reports = {} 
    
    def get_period_key(p):
        if isinstance(p, datetime.datetime):
            return p.date()
        return p

    # Merge CashBook
    for item in cashbook_data:
        period = get_period_key(item['period'])
        currency = item['currency']
        
        if period not in grouped_reports:
            grouped_reports[period] = {'LAK': {'income': 0, 'expense': 0}, 'THB': {'income': 0, 'expense': 0}, 'USD': {'income': 0, 'expense': 0}}
            
        if currency in grouped_reports[period]:
            grouped_reports[period][currency]['income'] += item['total_income'] or 0
            grouped_reports[period][currency]['expense'] += item['total_expense'] or 0

    # Merge Orders (Sales Income)
    for item in order_data:
        period = get_period_key(item['period'])
        if period not in grouped_reports:
             grouped_reports[period] = {'LAK': {'income': 0, 'expense': 0}, 'THB': {'income': 0, 'expense': 0}, 'USD': {'income': 0, 'expense': 0}}
        
        grouped_reports[period]['LAK']['income'] += item['income_lak'] or 0
        grouped_reports[period]['THB']['income'] += item['income_thb'] or 0
        grouped_reports[period]['USD']['income'] += item['income_usd'] or 0

    # Merge Stock Transactions (Stock Expenses - THB)
    for item in stock_data:
        period = get_period_key(item['period'])
        if period not in grouped_reports:
             grouped_reports[period] = {'LAK': {'income': 0, 'expense': 0}, 'THB': {'income': 0, 'expense': 0}, 'USD': {'income': 0, 'expense': 0}}
        
        # Add to THB Expense
        grouped_reports[period]['THB']['expense'] += item['expense_thb'] or 0

    # Calculate Balances & Format
    final_reports = []
    sorted_periods = sorted(grouped_reports.keys(), reverse=True)
    
    for period in sorted_periods:
        currencies = grouped_reports[period]
        # Calculate balance for each
        for curr in currencies:
            currencies[curr]['balance'] = currencies[curr]['income'] - currencies[curr]['expense']
            
        final_reports.append({
            'period': period,
            'data': currencies
        })
        
    return final_reports

@login_required
def report_view(request):
    """
    View for generating Accounting Reports (Daily, Weekly, Monthly, Yearly).
    Integrates data from:
    1. CashBook (Manual Income/Expense)
    2. Store Orders (Sales Income)
    3. Store Transactions (Stock/Purchase Expenses - Assumed THB)
    """
    
    report_type = request.GET.get('type', 'daily') # daily, weekly, monthly, yearly
    start_date_str = request.GET.get('start_date')
    end_date_str = request.GET.get('end_date')

    final_reports = _get_report_data(report_type, start_date_str, end_date_str)

    # --- Calculate Grand Totals for WhatsApp Summary ---
    grand_totals = {
        'LAK': {'income': 0, 'expense': 0, 'balance': 0},
        'THB': {'income': 0, 'expense': 0, 'balance': 0},
        'USD': {'income': 0, 'expense': 0, 'balance': 0},
    }

    for item in final_reports:
        data = item['data']
        for curr in ['LAK', 'THB', 'USD']:
            if curr in data:
                grand_totals[curr]['income'] += data[curr]['income']
                grand_totals[curr]['expense'] += data[curr]['expense']
                grand_totals[curr]['balance'] += data[curr]['balance']

    # --- Construct WhatsApp Text ---
    wa_start = start_date_str if start_date_str else "All Time"
    wa_end = end_date_str if end_date_str else ""
    wa_range = f"{wa_start} - {wa_end}" if wa_end else wa_start
    
    wa_title = f"📊 ບົດລາຍງານ ({report_type.title()})"
    lines = [f"{wa_title}", f"📅 ໄລຍະເວລາ: {wa_range}", "━━━━━━━━━━━━━━━━━━", ""]

    has_data = False
    
    # LAK
    if grand_totals['LAK']['income'] != 0 or grand_totals['LAK']['expense'] != 0:
        has_data = True
        lines.append(f"🟢 ₭ LAK (ກີບ)")
        lines.append(f"   ⬇️ ຮັບ: {grand_totals['LAK']['income']:,.0f}")
        lines.append(f"   ⬆️ ຈ່າຍ: {grand_totals['LAK']['expense']:,.0f}")
        lines.append(f"   💰 ຄົງເຫຼືອ: {grand_totals['LAK']['balance']:,.0f}")
        lines.append("")

    # THB
    if grand_totals['THB']['income'] != 0 or grand_totals['THB']['expense'] != 0:
        has_data = True
        lines.append(f"🟢 ฿ THB (ບາດ)")
        lines.append(f"   ⬇️ ຮັບ: {grand_totals['THB']['income']:,.0f}")
        lines.append(f"   ⬆️ ຈ່າຍ: {grand_totals['THB']['expense']:,.0f}")
        lines.append(f"   💰 ຄົງເຫຼືອ: {grand_totals['THB']['balance']:,.0f}")
        lines.append("")

    # USD
    if grand_totals['USD']['income'] != 0 or grand_totals['USD']['expense'] != 0:
        has_data = True
        lines.append(f"🟢 $ USD (ໂດລາ)")
        lines.append(f"   ⬇️ ຮັບ: {grand_totals['USD']['income']:,.0f}")
        lines.append(f"   ⬆️ ຈ່າຍ: {grand_totals['USD']['expense']:,.0f}")
        lines.append(f"   💰 ຄົງເຫຼືອ: {grand_totals['USD']['balance']:,.0f}")
        lines.append("")

    if not has_data:
        lines.append("❌ ບໍ່ມີການເຄື່ອນໄຫວໃນຊ່ວງເວລານີ້")

    from urllib.parse import quote
    wa_text = quote("\n".join(lines))

    context = {
        'report_type': report_type,
        'reports': final_reports,
        'start_date': start_date_str,
        'end_date': end_date_str,
        'whatsapp_text': wa_text,
    }
    
    return render(request, 'accounting/report.html', context)

@login_required
def export_report(request):
    report_type = request.GET.get('type', 'daily')
    start_date_str = request.GET.get('start_date')
    end_date_str = request.GET.get('end_date')
    export_format = request.GET.get('format', 'csv') # csv, excel, word, pdf
    
    reports = _get_report_data(report_type, start_date_str, end_date_str)
    timestamp = timezone.now().strftime('%Y%m%d')
    
    # ---------------- CSV EXPORT ----------------
    if export_format == 'csv':
        filename = f"report_{report_type}_{timestamp}.csv"
        response = HttpResponse(content_type='text/csv')
        response['Content-Disposition'] = f'attachment; filename="{filename}"'
        response.write(u'\ufeff'.encode('utf8')) # BOM

        writer = csv.writer(response)
        writer.writerow(['ເວລາ', 'ສະກຸນເງິນ', 'ລາຍຮັບ', 'ລາຍຈ່າຍ', 'ຍອດເງິນຄົງເຫຼືອ'])
        
        for report in reports:
            period = report['period']
            data = report['data']
            writer.writerow([period, 'LAK', data['LAK']['income'], data['LAK']['expense'], data['LAK']['balance']])
            writer.writerow([period, 'THB', data['THB']['income'], data['THB']['expense'], data['THB']['balance']])
            writer.writerow([period, 'USD', data['USD']['income'], data['USD']['expense'], data['USD']['balance']])
        return response

    # ---------------- EXCEL EXPORT ----------------
    elif export_format == 'excel':
        filename = f"report_{report_type}_{timestamp}.xlsx"
        response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
        response['Content-Disposition'] = f'attachment; filename="{filename}"'
        
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "Accounting Report"
        
        # Header
        headers = ['ເວລາ', 'ສະກຸນເງິນ', 'ລາຍຮັບ', 'ລາຍຈ່າຍ', 'ຍອດເງິນຄົງເຫຼືອ']
        ws.append(headers)
        
        # Style Header
        for cell in ws[1]:
            cell.font = Font(bold=True)
            cell.alignment = Alignment(horizontal='center')
            
        for report in reports:
            period = str(report['period'])
            data = report['data']
            
            ws.append([period, 'LAK', data['LAK']['income'], data['LAK']['expense'], data['LAK']['balance']])
            ws.append([period, 'THB', data['THB']['income'], data['THB']['expense'], data['THB']['balance']])
            ws.append([period, 'USD', data['USD']['income'], data['USD']['expense'], data['USD']['balance']])
            
        wb.save(response)
        return response

    # ---------------- WORD EXPORT ----------------
    elif export_format == 'word':
        filename = f"report_{report_type}_{timestamp}.docx"
        response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = f'attachment; filename="{filename}"'
        
        doc = Document()
        doc.add_heading(f'Accounting Report ({report_type.title()})', 0)
        
        table = doc.add_table(rows=1, cols=5)
        table.style = 'Table Grid'
        
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'ເວລາ'
        hdr_cells[1].text = 'ສະກຸນເງິນ'
        hdr_cells[2].text = 'ລາຍຮັບ'
        hdr_cells[3].text = 'ລາຍຈ່າຍ'
        hdr_cells[4].text = 'ຍອດເງິນຄົງເຫຼືອ'
        
        for report in reports:
            period = str(report['period'])
            data = report['data']
            
            # LAK
            row_lak = table.add_row().cells
            row_lak[0].text = period
            row_lak[1].text = 'LAK'
            row_lak[2].text = f"{data['LAK']['income']:,}"
            row_lak[3].text = f"{data['LAK']['expense']:,}"
            row_lak[4].text = f"{data['LAK']['balance']:,}"
            
            # THB
            row_thb = table.add_row().cells
            row_thb[0].text = period
            row_thb[1].text = 'THB'
            row_thb[2].text = f"{data['THB']['income']:,}"
            row_thb[3].text = f"{data['THB']['expense']:,}"
            row_thb[4].text = f"{data['THB']['balance']:,}"

            # USD
            row_usd = table.add_row().cells
            row_usd[0].text = period
            row_usd[1].text = 'USD'
            row_usd[2].text = f"{data['USD']['income']:,}"
            row_usd[3].text = f"{data['USD']['expense']:,}"
            row_usd[4].text = f"{data['USD']['balance']:,}"
            
        doc.save(response)
        return response

    # ---------------- PDF EXPORT ----------------
    elif export_format == 'pdf':
        filename = f"report_{report_type}_{timestamp}.pdf"
        response = HttpResponse(content_type='application/pdf')
        response['Content-Disposition'] = f'attachment; filename="{filename}"'
        
        buffer = BytesIO()
        doc = SimpleDocTemplate(
            buffer,
            pagesize=A4,
            rightMargin=50,
            leftMargin=50,
            topMargin=50,
            bottomMargin=50
        )
        elements = []
        
        # Register Font (Phetsarath OT)
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont
        import os
        from django.conf import settings
        from reportlab.lib.styles import ParagraphStyle

        font_path = os.path.join(settings.STATIC_ROOT, 'fonts', 'phetsarath_ot.ttf')
        # If running in development (runserver), staticfiles might not be collected to STATIC_ROOT
        if not os.path.exists(font_path):
             font_path = os.path.join(settings.BASE_DIR, 'autoparts', 'static', 'fonts', 'phetsarath_ot.ttf')
        
        try:
             pdfmetrics.registerFont(TTFont('Phetsarath OT', font_path))
             font_name = 'Phetsarath OT'
        except:
             font_name = 'Helvetica' # Fallback

        # Styles
        styles = getSampleStyleSheet()
        styles.add(ParagraphStyle(name='LaoTitle', parent=styles['Title'], fontName=font_name, fontSize=18, leading=22, alignment=1)) # Center Title
        styles.add(ParagraphStyle(name='LaoNormal', parent=styles['Normal'], fontName=font_name, fontSize=10, leading=14))
        
        # Title
        elements.append(Paragraph(f"ບົດລາຍງານ (Accounting Report) - {report_type.title()}", styles['LaoTitle']))
        elements.append(Spacer(1, 20))
        
        data_table = [['ເວລາ', 'ສະກຸນເງິນ', 'ລາຍຮັບ', 'ລາຍຈ່າຍ', 'ຍອດເງິນຄົງເຫຼືອ']]
        
        for report in reports:
            period = str(report['period'])
            for currency, data in report['data'].items():
                data_table.append([
                    period,
                    currency,
                    f"{data['income']:,.0f}",
                    f"{data['expense']:,.0f}",
                    f"{data['balance']:,.0f}"
                ])

        # Table Column Widths (A4 Width ~595pt. Margins 50+50=100. Available ~495pt)
        # Allocate: Time(100), Currency(50), Inc(115), Exp(115), Bal(115)
        col_widths = [100, 50, 115, 115, 115]

        # Table
        t = Table(data_table, colWidths=col_widths)
        t.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
            ('FONTNAME', (0, 0), (-1, -1), font_name),
            ('FONTSIZE', (0, 0), (-1, 0), 12),
            ('FONTSIZE', (0, 1), (-1, -1), 10),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('TOPPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), colors.white),
            ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
            ('ROWBACKGROUNDS', (1, 0), (-1, -1), [colors.whitesmoke, colors.white]), # Zebra Striping
        ]))
        
        elements.append(t)
        doc.build(elements)
        
        buffer.seek(0)
        response = HttpResponse(buffer, content_type='application/pdf')
        response['Content-Disposition'] = f'attachment; filename="accounting_report_{report_type}_{timestamp}.pdf"'
        return response

    return redirect('accounting:report')
